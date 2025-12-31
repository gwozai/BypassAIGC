"""
确定性渲染器：AST + reference.docx → output.docx
"""
from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass
from typing import Optional, Set

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm

from ..models.ast import (
    BibliographyBlock,
    DocumentAST,
    FigureBlock,
    HeadingBlock,
    ListBlock,
    ParagraphBlock,
    PageBreakBlock,
    SectionBreakBlock,
    TableBlock,
)
from ..models.stylespec import StyleSpec


def _align_to_docx(align: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]


class HeadingNumberer:
    """多级标题编号管理器。

    支持最多 6 级标题编号，格式如：1, 1.1, 1.1.1, 1.1.1.1, ...
    每当高级标题出现时，自动重置所有低级标题计数器。
    支持同步到已有编号，避免重复编号。
    """

    def __init__(self, max_level: int = 3, separator: str = "."):
        self._counters = [0] * 6  # 支持 6 级标题
        self._max_level = min(max_level, 6)
        self._separator = separator

    def reset(self) -> None:
        """重置所有计数器。"""
        self._counters = [0] * 6

    def sync_to_existing(self, existing_number: str, level: int) -> None:
        """同步计数器状态到已有编号。

        当 Markdown 中标题已有编号时，同步内部状态以避免重复。

        Args:
            existing_number: 已有的编号字符串，如 "1", "2.1", "3.1.2", "1．2"（含全角点号）
            level: 标题级别 (1-6)
        """
        if level < 1 or level > 6:
            return

        # 支持半角(.)和全角(．)点号分割
        parts = re.split(r'[.．]', existing_number)
        for i, part in enumerate(parts):
            if i < 6:
                try:
                    self._counters[i] = int(part)
                except ValueError:
                    pass

        # 重置所有更低级别的计数器
        for i in range(level, 6):
            self._counters[i] = 0

    def get_current_number(self, level: int) -> str:
        """获取指定级别的当前编号（不递增）。

        Args:
            level: 标题级别 (1-6)

        Returns:
            当前编号字符串
        """
        if level < 1 or level > self._max_level:
            return ""
        parts = [str(self._counters[i]) for i in range(level)]
        return self._separator.join(parts)

    def next_number(self, level: int) -> str:
        """获取指定级别的下一个编号。

        Args:
            level: 标题级别 (1-6)

        Returns:
            格式化的编号字符串，如 "1", "1.1", "1.1.1"
        """
        if level < 1 or level > self._max_level:
            return ""

        idx = level - 1

        # 递增当前级别计数器
        self._counters[idx] += 1

        # 重置所有更低级别的计数器
        for i in range(idx + 1, 6):
            self._counters[i] = 0

        # 构建编号字符串
        parts = [str(self._counters[i]) for i in range(level)]
        return self._separator.join(parts)


@dataclass
class RenderOptions:
    include_cover: bool = True
    include_toc: bool = True
    toc_title: str = "目 录"
    toc_levels: int = 3
    # 标题编号配置
    enable_heading_numbering: bool = True
    heading_numbering_max_level: int = 3
    heading_numbering_separator: str = "."
    heading_number_suffix: str = " "  # 编号与标题文本之间的分隔符


_FRONT_HEADINGS: Set[str] = {
    "摘要", "关键词", "关键字", "abstract", "key words", "keywords",
    "致谢", "谢辞", "参考文献", "references", "目录", "目 录",
}

_FRONT_ONLY_HEADINGS: Set[str] = {
    "摘要", "关键词", "关键字", "abstract", "key words", "keywords",
}

# 匹配已有编号的正则表达式：如 "1 标题", "1.2 标题", "1．2 标题"（含全角点号）
_EXISTING_HEADING_NUMBER_RE = re.compile(r"^(\d+(?:[.．]\d+)*)\s+(.+)$")


def _infer_level_from_number(number_str: str) -> int:
    """从编号字符串推断标题层级。

    Args:
        number_str: 编号字符串，如 "1", "1.1", "1.2.3", "1．1"（含全角点号）

    Returns:
        推断的层级数（1, 2, 3...）
    """
    if not number_str:
        return 1
    # 同时统计半角点号(.)和全角点号(．)
    dot_count = number_str.count('.') + number_str.count('．')
    return dot_count + 1


def _is_front_heading(text: str) -> bool:
    """检查是否为前置标题（不区分大小写）。"""
    return text.lower() in _FRONT_HEADINGS or text in _FRONT_HEADINGS


def _is_front_only_heading(text: str) -> bool:
    """检查是否为仅前置标题（不区分大小写）。"""
    return text.lower() in _FRONT_ONLY_HEADINGS or text in _FRONT_ONLY_HEADINGS


def _apply_page_numbering_ooxml(docx_bytes: bytes, spec: StyleSpec) -> bytes:
    """Set section-based page numbering format/start using OOXML."""
    pn = spec.page_numbering
    if not pn or not pn.enabled:
        return docx_bytes
    from ..utils.ooxml import DocxPackage
    from lxml import etree

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NSMAP = {"w": W_NS}

    def _qn(tag: str) -> str:
        pref, local = tag.split(":")
        return f"{{{NSMAP[pref]}}}{local}"

    fmt_map = {
        "decimal": "decimal",
        "romanUpper": "upperRoman",
        "romanLower": "lowerRoman",
    }

    pkg = DocxPackage.from_bytes(docx_bytes)
    root = pkg.read_xml("word/document.xml")
    body = root.find("w:body", namespaces=NSMAP)
    if body is None:
        return docx_bytes

    sect_prs = body.findall(".//w:sectPr", namespaces=NSMAP)
    seen = set()
    ordered = []
    for s in sect_prs:
        sid = id(s)
        if sid in seen:
            continue
        seen.add(sid)
        ordered.append(s)
    if not ordered:
        return docx_bytes

    def _set_pgnum(sectPr, fmt: str, start: int):
        pg = sectPr.find("w:pgNumType", namespaces=NSMAP)
        if pg is None:
            pg = etree.SubElement(sectPr, _qn("w:pgNumType"))
        pg.set(_qn("w:fmt"), fmt_map.get(fmt, "decimal"))
        pg.set(_qn("w:start"), str(int(start)))

    if len(ordered) == 1:
        _set_pgnum(ordered[0], pn.main_format, pn.main_start)
    else:
        _set_pgnum(ordered[0], pn.front_format, pn.front_start)
        _set_pgnum(ordered[-1], pn.main_format, pn.main_start)

    pkg.write_xml("word/document.xml", root)
    return pkg.to_bytes()


def _insert_toc_paragraph(doc: Document, title: str, front_style: str, max_level: int):
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph(title)
    p.style = doc.styles[front_style]
    p2 = doc.add_paragraph()
    run = p2.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\o "1-{max_level}" \\h \\z \\u')
    run._r.append(fld)
    run2 = p2.add_run("（在 Word 中右键目录 → 更新域）")


def _detect_heading_level_offset(ast: DocumentAST) -> int:
    """检测标题级别偏移量。

    如果文档中最小的标题级别不是 1（例如用户用 ## 作为一级标题），
    计算需要减去的偏移量，使其归一化为从 1 开始。

    Args:
        ast: 文档 AST

    Returns:
        偏移量（0 表示无需调整，1 表示所有级别减 1，以此类推）
    """
    min_level = None
    for block in ast.blocks:
        if isinstance(block, HeadingBlock):
            # 跳过前置标题（摘要、Abstract 等通常是单独设置的）
            heading_text = block.text.strip().lower()
            if heading_text in _FRONT_HEADINGS:
                continue
            if min_level is None or block.level < min_level:
                min_level = block.level

    # 如果没有正文标题或最小级别已是 1，无需调整
    if min_level is None or min_level <= 1:
        return 0

    return min_level - 1


def render_docx(
    ast: DocumentAST,
    spec: StyleSpec,
    reference_docx_bytes: bytes,
    options: Optional[RenderOptions] = None,
) -> bytes:
    options = options or RenderOptions()
    doc = Document(io.BytesIO(reference_docx_bytes))

    section = doc.sections[0]
    section.top_margin = Mm(spec.page.margins_mm.top)
    section.bottom_margin = Mm(spec.page.margins_mm.bottom)
    section.left_margin = Mm(spec.page.margins_mm.left)
    section.right_margin = Mm(spec.page.margins_mm.right)
    section.gutter = Mm(spec.page.margins_mm.binding)
    section.header_distance = Mm(spec.page.header_mm)
    section.footer_distance = Mm(spec.page.footer_mm)

    # clear initial dummy paragraph if empty
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Cover
    if options.include_cover:
        _render_cover(doc, ast)
        doc.add_page_break()

    # TOC
    if options.include_toc:
        _insert_toc_paragraph(doc, options.toc_title, "FrontHeading", spec.structure.toc_max_level)
        doc.add_page_break()

    need_page_numbering = bool(spec.page_numbering and spec.page_numbering.enabled)
    main_section_inserted = False

    current_section = None
    abstract_prefixed = False
    keywords_prefixed = False
    fig_counter = 0
    table_counter = 0

    # 检测标题级别偏移（支持 ## 作为一级标题等情况）
    heading_level_offset = _detect_heading_level_offset(ast)

    # 初始化标题编号器
    heading_numberer = HeadingNumberer(
        max_level=options.heading_numbering_max_level,
        separator=options.heading_numbering_separator,
    ) if options.enable_heading_numbering else None

    for block in ast.blocks:
        if isinstance(block, HeadingBlock):
            heading_text = block.text.strip()

            # 检测已有的编号前缀（如 "1 标题", "1.2 标题"）
            existing_match = _EXISTING_HEADING_NUMBER_RE.match(heading_text)
            existing_number = None
            if existing_match:
                existing_number = existing_match.group(1)  # 保存已有编号
                heading_text = existing_match.group(2)     # 提取纯文本

            if heading_text in {"摘要"}:
                current_section = "cn_abstract"
            elif heading_text in {"关键词", "关键字"}:
                current_section = "cn_keywords"
            elif heading_text.lower() == "abstract":
                current_section = "en_abstract"
            elif heading_text.lower() in {"key words", "keywords"}:
                current_section = "en_keywords"
            elif heading_text.lower() in {"参考文献", "references"}:
                current_section = "references"
            else:
                current_section = "body"

            if (
                need_page_numbering
                and not main_section_inserted
                and not _is_front_only_heading(heading_text)
                and len(doc.paragraphs) > 0
            ):
                doc.add_section(WD_SECTION.NEW_PAGE)
                main_section_inserted = True

            if _is_front_heading(heading_text):
                style_id = "FrontHeading"
                display_text = heading_text
            else:
                # 确定有效层级：优先使用编号推断的层级，否则使用 Markdown 语法层级
                if existing_number:
                    # 从编号推断层级（如 "1.1" -> 层级 2）
                    effective_level = _infer_level_from_number(existing_number)
                else:
                    # 应用级别偏移（支持 ## 作为一级标题等情况）
                    effective_level = max(1, block.level - heading_level_offset)

                # 限制最大层级为 3（模板通常只支持 H1-H3）
                effective_level = min(effective_level, 3)

                if effective_level == 1:
                    style_id = "H1"
                elif effective_level == 2:
                    style_id = "H2"
                elif effective_level == 3:
                    style_id = "H3"
                else:
                    style_id = "H3"

                # 为非前置标题处理编号
                if heading_numberer and effective_level <= options.heading_numbering_max_level:
                    if existing_number:
                        # Markdown 已有编号：同步状态并直接使用原编号
                        heading_numberer.sync_to_existing(existing_number, effective_level)
                        display_text = f"{existing_number}{options.heading_number_suffix}{heading_text}"
                    else:
                        # Markdown 无编号：自动生成新编号
                        number = heading_numberer.next_number(effective_level)
                        display_text = f"{number}{options.heading_number_suffix}{heading_text}"
                else:
                    display_text = heading_text

            p = doc.add_paragraph(display_text)
            if style_id in doc.styles:
                p.style = doc.styles[style_id]
            else:
                p.style = doc.styles["Body"]
            continue

        if isinstance(block, ParagraphBlock):
            txt = block.text
            if txt is None and block.inlines:
                txt = "".join(i.text for i in block.inlines)
            txt = (txt or "").strip()
            if not txt:
                continue

            if spec.auto_prefix_abstract_keywords:
                if current_section == "cn_abstract" and not abstract_prefixed:
                    if not txt.startswith("摘要："):
                        txt = "摘要：" + txt
                    abstract_prefixed = True
                elif current_section == "en_abstract" and not abstract_prefixed:
                    if not txt.lower().startswith("abstract:"):
                        txt = "Abstract: " + txt
                    abstract_prefixed = True
                elif current_section in {"cn_keywords", "en_keywords"} and not keywords_prefixed:
                    if current_section == "cn_keywords" and not txt.startswith(("关键词：", "关键字：")):
                        txt = "关键词：" + _normalize_cn_keywords(txt)
                    elif current_section == "en_keywords" and not txt.lower().startswith(("key words:", "keywords:")):
                        txt = "Key words: " + _normalize_en_keywords(txt)
                    keywords_prefixed = True

            style_id = "Body"
            if current_section in {"cn_abstract", "en_abstract"}:
                style_id = "AbstractBody"
            elif current_section in {"cn_keywords", "en_keywords"}:
                style_id = "KeywordsBody"
            elif current_section == "references":
                style_id = "Reference"

            p = doc.add_paragraph(txt)
            if style_id in doc.styles:
                p.style = doc.styles[style_id]
            continue

        if isinstance(block, ListBlock):
            style_name = "ListNumber" if block.ordered else "ListBullet"
            use_style = style_name in doc.styles
            for idx, item in enumerate(block.items, start=1):
                txt = "".join(i.text for i in item.inlines).strip()
                if not txt:
                    continue
                if use_style:
                    p = doc.add_paragraph(txt)
                    p.style = doc.styles[style_name]
                else:
                    prefix = f"{idx}. " if block.ordered else "• "
                    p = doc.add_paragraph(prefix + txt)
                    p.style = doc.styles["Body"]
            continue

        if isinstance(block, TableBlock):
            if block.caption:
                caption = block.caption.strip()
                if spec.auto_number_figures_tables and not re.match(r"^表\d+", caption):
                    table_counter += 1
                    caption = f"表{table_counter} {caption}"
                pcap = doc.add_paragraph(caption)
                if "TableTitle" in doc.styles:
                    pcap.style = doc.styles["TableTitle"]
            elif spec.auto_number_figures_tables:
                table_counter += 1
                pcap = doc.add_paragraph(f"表{table_counter}")
                if "TableTitle" in doc.styles:
                    pcap.style = doc.styles["TableTitle"]
            if not block.rows:
                continue
            cols = max(len(r) for r in block.rows)
            table = doc.add_table(rows=len(block.rows), cols=cols)
            for r_i, row in enumerate(block.rows):
                for c_i in range(cols):
                    cell = table.cell(r_i, c_i)
                    cell.text = row[c_i] if c_i < len(row) else ""
                    for p in cell.paragraphs:
                        if "TableText" in doc.styles:
                            p.style = doc.styles["TableText"]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _apply_three_line_table(table)
            continue

        if isinstance(block, FigureBlock):
            if spec.auto_number_figures_tables and block.caption and not re.match(r"^图\d+", block.caption.strip()):
                fig_counter += 1
                caption = f"图{fig_counter} {block.caption.strip()}"
            else:
                caption = block.caption
            if os.path.exists(block.path):
                doc.add_picture(block.path)
            else:
                p = doc.add_paragraph(f"[图片占位：{block.path}]")
                p.style = doc.styles["Body"]
            if caption:
                pcap = doc.add_paragraph(caption)
                if "FigureCaption" in doc.styles:
                    pcap.style = doc.styles["FigureCaption"]
            continue

        if isinstance(block, PageBreakBlock):
            doc.add_page_break()
            continue

        if isinstance(block, SectionBreakBlock):
            doc.add_section(WD_SECTION.NEW_PAGE)
            continue

        if isinstance(block, BibliographyBlock):
            for it in block.items:
                p = doc.add_paragraph(it)
                if "Reference" in doc.styles:
                    p.style = doc.styles["Reference"]
            continue

    out = io.BytesIO()
    doc.save(out)
    data = out.getvalue()
    if need_page_numbering:
        _ensure_footer_page_numbers(doc, spec)
        out = io.BytesIO()
        doc.save(out)
        data = out.getvalue()
        data = _apply_page_numbering_ooxml(data, spec)
    return data


def _render_cover(doc: Document, ast: DocumentAST) -> None:
    if ast.meta.title_cn:
        p = doc.add_paragraph(ast.meta.title_cn)
        if "TitleCN" in doc.styles:
            p.style = doc.styles["TitleCN"]
    if ast.meta.title_en:
        p = doc.add_paragraph(ast.meta.title_en)
        if "TitleEN" in doc.styles:
            p.style = doc.styles["TitleEN"]

    meta_parts = []
    if ast.meta.major:
        meta_parts.append(f"专业：{ast.meta.major}")
    if ast.meta.author:
        meta_parts.append(f"学生：{ast.meta.author}")
    if ast.meta.tutor:
        meta_parts.append(f"指导教师：{ast.meta.tutor}")
    for line in meta_parts:
        p = doc.add_paragraph(line)
        if "MetaLine" in doc.styles:
            p.style = doc.styles["MetaLine"]


def _ensure_footer_page_numbers(doc: Document, spec: StyleSpec) -> None:
    pn = spec.page_numbering
    if not pn or not pn.enabled or not pn.show_in_footer:
        return

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for r in list(p.runs):
            try:
                p._p.remove(r._r)
            except Exception:
                pass
        if "PageNumber" in doc.styles:
            p.style = doc.styles["PageNumber"]
        p.alignment = _align_to_docx(pn.footer_alignment)
        run = p.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.append(fld)


def _normalize_cn_keywords(txt: str) -> str:
    parts = [p for p in re.split(r"[，,;；\s]+", txt) if p]
    return "　".join(parts)


def _normalize_en_keywords(txt: str) -> str:
    parts = [p.strip() for p in re.split(r"[;；,，]+", txt) if p.strip()]
    return "; ".join(parts)


def _apply_three_line_table(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    def _border(tag: str, val: str, sz: int):
        el = tblBorders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tblBorders.append(el)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    _border("top", "single", 12)
    _border("bottom", "single", 12)
    _border("insideH", "single", 6)
    _border("left", "nil", 0)
    _border("right", "nil", 0)
    _border("insideV", "nil", 0)
